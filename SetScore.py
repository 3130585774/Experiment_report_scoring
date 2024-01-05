import os

import openpyxl
from docx import Document
from docx.shared import Pt

from ChanageName import name_dict, find_key_value_pairs

className = "1班"


def init_excel():
    workbook = openpyxl.Workbook()
    sheet = workbook.active

    index_list = ["id", "name"]

    # for lesson in list_subdirectories(className):
    #     lesson_numbers = list_subdirectories(className + "/" + lesson)
    #     for little_lesson in lesson_numbers:
    #         index_list.append(little_lesson)
    for i in range(1, 8):
        index_list.append(f"go语言实验报告{i}")

    print(index_list)
    sheet.append(index_list)

    for key, value in name_dict.items():
        sheet.append([key, value])
    # if file ext
    if os.path.exists(f"{className}.xlsx"):
        print(f"{className}.xlsx 存在")
    else:
        workbook.save(f"{className}.xlsx")


def setScore(_fileName, num: int):
    """设置分数"""
    try:
        # print(_fileName)
        doc = Document(_fileName)
        # 获取表格对象
        table = doc.tables[0]
        # 获取要插入文本的单元格
        cell = table.cell(3, 6)

        cell.text = f"成绩：{num}"
        cell.paragraphs[0].runs[0].font.name = '宋体'  # 设置字体为宋体
        cell.paragraphs[0].runs[0].font.size = Pt(12)  # 设置字体大小为12磅

        # 保存文档
        doc.save(_fileName)
    except Exception:
        print(_fileName, "成绩为", num)
        # try:
        #     os.startfile(f"./{_fileName}")
        # except Exception as e:
        #     print(_fileName, "打开失败", e)


def getScoreDict(file):
    workbook = openpyxl.load_workbook(file)

    sheet = workbook.active

    all_rows = sheet.iter_rows(values_only=True)

    columns = next(all_rows)

    data_dict = {column: [] for column in columns}

    for row in all_rows:
        for col_index, value in enumerate(row):
            data_dict[columns[col_index]].append(value)

    workbook.close()
    return data_dict


def getScoreById(_id, subject, score_dict):
    try:
        score = score_dict[subject][score_dict['id'].index(int(_id))]
        # print(score)
    except Exception:
        print(f"Error:{_id}")
        score = 0
    return int(score)


def getScoreByFileName(data_list, little, fileName):
    key = find_key_value_pairs(fileName, name_dict)
    # print(key[0], end="")
    score = getScoreById(key[0], little, data_list)
    return score
