from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook
import os


def putNumber(_fileName, name, _classNumber, num: int):
    try:
        # print(_fileName)
        doc = Document(_fileName)
        # 获取表格对象
        table = doc.tables[0]
        # 获取要插入文本的单元格
        cell = table.cell(3, 6)
        # 插入带有样式的文本
        # 替换单元格中的文本
        cell.text = f"成绩：{num}"
        cell.paragraphs[0].runs[0].font.name = '宋体'  # 设置字体为宋体
        cell.paragraphs[0].runs[0].font.size = Pt(12)  # 设置字体大小为12磅
        # new_paragraph = cell.add_paragraph()
        # new_run = new_paragraph.add_run(f"成绩：{num}")
        # new_run.font.name = '宋体'  # 设置字体为宋体
        # new_run.font.size = Pt(12)  # 设置字体大小为12磅
        # 保存文档
        doc.save(f"{_fileName[:19]}+{name}_数据结构综合设计（{_classNumber}班）.docx")
        if _fileName != f"{_fileName[:19]}+{name}_数据结构综合设计（{_classNumber}班）.docx":
            os.remove("E:\docTest" + _fileName[1:])
    except Exception:
        print(_fileName[6:], "成绩为", num)
        os.startfile("E:\docTest" + _fileName[1:])


def getFolderFileNameList(folder_path):
    # 指定文件夹路径
    # folder_path = './'
    fileNameList = []
    # 获取文件夹中的所有文件名
    file_names = os.listdir(folder_path)

    # 输出所有文件名
    for file_name in file_names:
        if os.path.isfile(os.path.join(folder_path, file_name)):
            # print(file_name)
            fileNameList.append(file_name)
    return fileNameList


# 打开 Excel 文件
workbook = load_workbook(filename='三个班实验报告成绩.xlsx')
for classNumber in range(1, 4):
    print(classNumber, "班")
    students = []
    worksheet = workbook[f'{classNumber}班']
    # 遍历工作表中的每一行
    for row in worksheet.iter_rows(min_row=2, max_row=50, values_only=True):

        # 输出每一行所包含的数据
        row_data = [cell if cell is not None else '' for cell in row]
        if row_data[0]:
            # print(row_data)
            students.append(row_data)
    for time in range(1, 4):
        print(time, "次")
        fileList = getFolderFileNameList(f"./{classNumber}/{time}")
        for student in students:
            for fileName in fileList:
                if student[0] in fileName:
                    putNumber(f"./{classNumber}/{time}/{fileName}", student[0], classNumber, student[time])

# 关闭 Excel 文件
workbook.close()
